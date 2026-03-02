from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import HttpResponse
from django.template.loader import render_to_string
from django.db.models import Sum, Count, Q, F
from django.utils import timezone
from datetime import timedelta
from inventory.models import Component, Category, StockMovement
from orders.models import Order, OrderItem
from decimal import Decimal
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

# WeasyPrint импортируется условно, т.к. требует дополнительных библиотек на Windows
try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError):
    WEASYPRINT_AVAILABLE = False


@login_required
def reports_home(request):
    """Главная страница отчетов."""
    if not (request.user.is_worker or request.user.is_superuser):
        messages.error(request, 'У вас нет доступа к этой странице.')
        from django.shortcuts import redirect
        return redirect('core:dashboard')
    
    return render(request, 'reports/home.html')


@login_required
def inventory_report(request):
    """Отчет по складу."""
    if not (request.user.is_worker or request.user.is_superuser):
        messages.error(request, 'У вас нет доступа к этой странице.')
        from django.shortcuts import redirect
        return redirect('core:dashboard')
    
    # Фильтры по датам
    date_from = request.GET.get('date_from')
    date_to = request.GET.get('date_to')
    
    components = Component.objects.select_related('category').all()
    
    # Фильтрация по дате создания компонентов
    if date_from:
        components = components.filter(created_at__gte=date_from)
    if date_to:
        components = components.filter(created_at__lte=date_to)
    
    # Добавляем вычисленные суммы для каждого компонента
    components_with_totals = []
    for c in components:
        components_with_totals.append({
            'obj': c,
            'total': c.price * c.quantity
        })
    
    # Статистика
    total_components = components.count()
    low_stock_components = components.filter(quantity__lte=F('min_quantity')).count()
    out_of_stock = components.filter(quantity=0).count()
    total_value = sum(c.price * c.quantity for c in components)
    
    # Группировка по категориям
    categories_stats = Category.objects.annotate(
        total_items=Count('components'),
        total_quantity=Sum('components__quantity')
    )
    
    context = {
        'components': components_with_totals,
        'total_components': total_components,
        'low_stock_components': low_stock_components,
        'out_of_stock': out_of_stock,
        'total_value': total_value,
        'categories_stats': categories_stats,
        'date_from': date_from,
        'date_to': date_to,
    }
    
    return render(request, 'reports/inventory_report.html', context)


@login_required
def inventory_report_pdf(request):
    """PDF отчет по складу."""
    if not (request.user.is_worker or request.user.is_superuser):
        messages.error(request, 'У вас нет доступа к этой функции.')
        from django.shortcuts import redirect
        return redirect('core:dashboard')
    
    # TODO: Implement PDF generation using WeasyPrint or ReportLab
    components = Component.objects.select_related('category').all()
    
    html = render_to_string('reports/inventory_report_pdf.html', {
        'components': components,
    })
    
    return HttpResponse(html, content_type='text/html')


@login_required
def orders_report(request):
    """Отчет по заявкам."""
    if not (request.user.is_worker or request.user.is_superuser):
        messages.error(request, 'У вас нет доступа к этой странице.')
        from django.shortcuts import redirect
        return redirect('core:dashboard')
    
    # Фильтры по датам
    date_from = request.GET.get('date_from')
    date_to = request.GET.get('date_to')
    
    orders = Order.objects.select_related('client', 'assigned_to').all()
    
    if date_from:
        orders = orders.filter(created_at__gte=date_from)
    if date_to:
        orders = orders.filter(created_at__lte=date_to)
    
    # Статистика
    total_orders = orders.count()
    new_orders = orders.filter(status='new').count()
    in_progress_orders = orders.filter(status='in_progress').count()
    completed_orders = orders.filter(status='completed').count()
    cancelled_orders = orders.filter(status='cancelled').count()
    
    # Статистика по типам
    service_orders = orders.filter(order_type='service').count()
    component_orders = orders.filter(order_type='components').count()
    
    context = {
        'orders': orders,
        'total_orders': total_orders,
        'new_orders': new_orders,
        'in_progress_orders': in_progress_orders,
        'completed_orders': completed_orders,
        'cancelled_orders': cancelled_orders,
        'service_orders': service_orders,
        'component_orders': component_orders,
        'date_from': date_from,
        'date_to': date_to,
    }
    
    return render(request, 'reports/orders_report.html', context)


@login_required
def orders_report_pdf(request):
    """PDF отчет по заявкам."""
    if not (request.user.is_worker or request.user.is_superuser):
        messages.error(request, 'У вас нет доступа к этой функции.')
        from django.shortcuts import redirect
        return redirect('core:dashboard')
    
    # TODO: Implement PDF generation
    orders = Order.objects.select_related('client', 'assigned_to').all()
    
    html = render_to_string('reports/orders_report_pdf.html', {
        'orders': orders,
    })
    
    return HttpResponse(html, content_type='text/html')


@login_required
def financial_report(request):
    """Финансовый отчет."""
    if not (request.user.is_worker or request.user.is_superuser):
        messages.error(request, 'У вас нет доступа к этой странице.')
        from django.shortcuts import redirect
        return redirect('core:dashboard')
    
    # Фильтры по датам
    date_from = request.GET.get('date_from')
    date_to = request.GET.get('date_to')
    
    orders = Order.objects.filter(status='completed')
    
    if date_from:
        orders = orders.filter(completed_at__gte=date_from)
    if date_to:
        orders = orders.filter(completed_at__lte=date_to)
    
    # Статистика
    total_revenue = orders.aggregate(total=Sum('total_amount'))['total'] or Decimal('0.00')
    total_completed = orders.count()
    
    # Средний чек
    avg_order_value = total_revenue / total_completed if total_completed > 0 else Decimal('0.00')
    
    # Доход по типам заявок
    service_revenue = orders.filter(order_type='service').aggregate(
        total=Sum('total_amount'))['total'] or Decimal('0.00')
    components_revenue = orders.filter(order_type='components').aggregate(
        total=Sum('total_amount'))['total'] or Decimal('0.00')
    
    context = {
        'orders': orders,
        'total_revenue': total_revenue,
        'total_completed': total_completed,
        'avg_order_value': avg_order_value,
        'service_revenue': service_revenue,
        'components_revenue': components_revenue,
        'date_from': date_from,
        'date_to': date_to,
    }
    
    return render(request, 'reports/financial_report.html', context)


@login_required
def financial_report_pdf(request):
    """PDF финансового отчета."""
    if not (request.user.is_worker or request.user.is_superuser):
        messages.error(request, 'У вас нет доступа к этой функции.')
        from django.shortcuts import redirect
        return redirect('core:dashboard')
    
    # Получаем параметры фильтрации дат
    date_from = request.GET.get('date_from')
    date_to = request.GET.get('date_to')
    
    if date_from:
        date_from = timezone.datetime.strptime(date_from, '%Y-%m-%d').date()
    if date_to:
        date_to = timezone.datetime.strptime(date_to, '%Y-%m-%d').date()
    
    # Фильтруем заявки
    orders = Order.objects.filter(status='completed')
    
    if date_from:
        orders = orders.filter(completed_at__gte=date_from)
    if date_to:
        orders = orders.filter(completed_at__lte=date_to)
    
    # Статистика
    total_revenue = orders.aggregate(total=Sum('total_amount'))['total'] or Decimal('0.00')
    total_completed = orders.count()
    avg_order_value = total_revenue / total_completed if total_completed > 0 else Decimal('0.00')
    
    # Доход по типам заявок
    service_revenue = orders.filter(order_type='service').aggregate(
        total=Sum('total_amount'))['total'] or Decimal('0.00')
    components_revenue = orders.filter(order_type='components').aggregate(
        total=Sum('total_amount'))['total'] or Decimal('0.00')
    
    html = render_to_string('reports/financial_report_pdf.html', {
        'orders': orders,
        'total_revenue': total_revenue,
        'total_completed': total_completed,
        'avg_order_value': avg_order_value,
        'service_revenue': service_revenue,
        'components_revenue': components_revenue,
        'date_from': date_from,
        'date_to': date_to,
        'now': timezone.now(),
    })
    
    return HttpResponse(html, content_type='text/html')


# ==================== EXCEL EXPORTS ====================

@login_required
def inventory_report_excel(request):
    """Экспорт отчета по складу в Excel."""
    if not (request.user.is_worker or request.user.is_superuser):
        messages.error(request, 'У вас нет доступа к этой функции.')
        from django.shortcuts import redirect
        return redirect('core:dashboard')
    
    # Получаем параметры фильтрации
    date_from = request.GET.get('date_from')
    date_to = request.GET.get('date_to')
    
    # Создаем книгу Excel
    wb = Workbook()
    ws = wb.active
    ws.title = "Отчет по складу"
    
    # Заголовок
    ws.merge_cells('A1:G1')
    title_cell = ws['A1']
    title_cell.value = "ОТЧЕТ ПО СКЛАДУ"
    title_cell.font = Font(size=16, bold=True)
    title_cell.alignment = Alignment(horizontal='center', vertical='center')
    title_cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    title_cell.font = Font(size=16, bold=True, color="FFFFFF")
    
    # Период отчета
    if date_from or date_to:
        ws.append([''])
        period_text = 'Период: '
        if date_from:
            period_text += f'с {date_from} '
        if date_to:
            period_text += f'по {date_to}'
        ws.append([period_text])
    
    # Заголовки столбцов
    headers = ['Артикул', 'Название', 'Категория', 'Количество', 'Мин. количество', 'Цена', 'Сумма']
    ws.append([''])  # Пустая строка
    ws.append(headers)
    
    # Стиль заголовков
    header_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    for cell in ws[3]:
        cell.font = Font(bold=True)
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center')
        cell.border = border
    
    # Данные
    components = Component.objects.select_related('category').all()
    
    # Применяем фильтрацию по датам
    if date_from:
        components = components.filter(created_at__gte=date_from)
    if date_to:
        components = components.filter(created_at__lte=date_to)
    
    for component in components:
        total = component.price * component.quantity
        row = [
            component.article_number,
            component.name,
            component.category.name if component.category else '-',
            component.quantity,
            component.min_quantity,
            float(component.price),
            float(total)
        ]
        ws.append(row)
        
        # Применяем стиль к ячейкам
        for cell in ws[ws.max_row]:
            cell.border = border
            if cell.column in [4, 5]:  # Количество
                cell.alignment = Alignment(horizontal='center')
            elif cell.column in [6, 7]:  # Цены
                cell.number_format = '#,##0.00 ₽'
    
    # Итоги
    ws.append([''])
    total_row = ws.max_row + 1
    ws[f'A{total_row}'] = 'ИТОГО:'
    ws[f'A{total_row}'].font = Font(bold=True)
    ws[f'G{total_row}'] = f'=SUM(G4:G{total_row-2})'
    ws[f'G{total_row}'].font = Font(bold=True)
    ws[f'G{total_row}'].number_format = '#,##0.00 ₽'
    
    # Автоширина столбцов
    for idx, column in enumerate(ws.columns, 1):
        max_length = 0
        column_letter = get_column_letter(idx)
        for cell in column:
            try:
                if cell.value and len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column_letter].width = adjusted_width
    
    # Сохраняем в память
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    
    # Возвращаем файл
    response = HttpResponse(
        output.read(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename="inventory_report_{timezone.now().strftime("%Y%m%d")}.xlsx"'
    return response


@login_required
def orders_report_excel(request):
    """Экспорт отчета по заявкам в Excel."""
    if not (request.user.is_worker or request.user.is_superuser):
        messages.error(request, 'У вас нет доступа к этой функции.')
        from django.shortcuts import redirect
        return redirect('core:dashboard')
    
    # Получаем параметры фильтрации
    date_from = request.GET.get('date_from')
    date_to = request.GET.get('date_to')
    
    # Создаем книгу Excel
    wb = Workbook()
    ws = wb.active
    ws.title = "Отчет по заявкам"
    
    # Заголовок
    ws.merge_cells('A1:H1')
    title_cell = ws['A1']
    title_cell.value = "ОТЧЕТ ПО ЗАЯВКАМ"
    title_cell.font = Font(size=16, bold=True)
    title_cell.alignment = Alignment(horizontal='center', vertical='center')
    title_cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    title_cell.font = Font(size=16, bold=True, color="FFFFFF")
    
    # Период отчета
    if date_from or date_to:
        ws.append([''])
        period_text = 'Период: '
        if date_from:
            period_text += f'с {date_from} '
        if date_to:
            period_text += f'по {date_to}'
        ws.append([period_text])
    
    # Заголовки столбцов
    headers = ['№ Заявки', 'Дата', 'Клиент', 'Тип', 'Устройство', 'Статус', 'Исполнитель', 'Сумма']
    ws.append([''])
    ws.append(headers)
    
    # Стиль заголовков
    header_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    for cell in ws[3]:
        cell.font = Font(bold=True)
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center')
        cell.border = border
    
    # Данные
    orders = Order.objects.select_related('client', 'assigned_to').all()
    
    # Применяем фильтрацию по датам
    if date_from:
        orders = orders.filter(created_at__gte=date_from)
    if date_to:
        orders = orders.filter(created_at__lte=date_to)
    
    status_dict = dict(Order.STATUS_CHOICES)
    type_dict = dict(Order.ORDER_TYPE_CHOICES)
    
    for order in orders:
        row = [
            order.id,
            order.created_at.strftime('%d.%m.%Y %H:%M'),
            order.client.get_full_name() or order.client.username,
            type_dict.get(order.order_type, order.order_type),
            order.device_type or '-',
            status_dict.get(order.status, order.status),
            order.assigned_to.get_full_name() if order.assigned_to else 'Не назначен',
            float(order.total_amount) if order.total_amount else 0
        ]
        ws.append(row)
        
        # Применяем стиль
        for cell in ws[ws.max_row]:
            cell.border = border
            if cell.column == 8:  # Сумма
                cell.number_format = '#,##0.00 ₽'
    
    # Итоги
    ws.append([''])
    total_row = ws.max_row + 1
    ws[f'A{total_row}'] = 'ИТОГО:'
    ws[f'A{total_row}'].font = Font(bold=True)
    ws[f'H{total_row}'] = f'=SUM(H4:H{total_row-2})'
    ws[f'H{total_row}'].font = Font(bold=True)
    ws[f'H{total_row}'].number_format = '#,##0.00 ₽'
    
    # Автоширина столбцов
    for idx, column in enumerate(ws.columns, 1):
        max_length = 0
        column_letter = get_column_letter(idx)
        for cell in column:
            try:
                if cell.value and len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column_letter].width = adjusted_width
    
    # Сохраняем в память
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    
    # Возвращаем файл
    response = HttpResponse(
        output.read(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename="orders_report_{timezone.now().strftime("%Y%m%d")}.xlsx"'
    return response


# ==================== WORD EXPORTS ====================

@login_required
def inventory_report_word(request):
    """Экспорт отчета по складу в Word."""
    if not (request.user.is_worker or request.user.is_superuser):
        messages.error(request, 'У вас нет доступа к этой функции.')
        from django.shortcuts import redirect
        return redirect('core:dashboard')
    
    # Получаем параметры фильтрации
    date_from = request.GET.get('date_from')
    date_to = request.GET.get('date_to')
    
    # Создаем документ
    doc = Document()
    
    # Шапка с информацией о компании
    company_heading = doc.add_paragraph()
    company_run = company_heading.add_run('ООО НПФ «Роникс-Л»')
    company_run.bold = True
    company_run.font.size = Pt(14)
    company_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    company_details = doc.add_paragraph()
    details_run = company_details.add_run(
        'ОГРН: 1021601978177 | ИНН: 1649003743 / КПП: 164901001\n'
        '423258, Республика Татарстан, Лениногорский район,\n'
        'город Лениногорск, Крупской ул, д. 4, помещ. 3'
    )
    details_run.font.size = Pt(9)
    company_details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('_' * 80)
    
    # Заголовок отчета
    title = doc.add_heading('ОТЧЕТ ПО СОСТОЯНИЮ СКЛАДСКИХ ЗАПАСОВ', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Дата формирования
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f'Дата формирования: {timezone.now().strftime("%d.%m.%Y %H:%M")}')
    date_run.italic = True
    date_run.font.size = Pt(10)
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()
    
    # Назначение отчета
    doc.add_heading('1. НАЗНАЧЕНИЕ ОТЧЕТА', level=1)
    purpose_text = (
        'Настоящий отчет предназначен для анализа текущего состояния товарных запасов '
        'на складе предприятия. Отчет содержит информацию о номенклатуре товаров, '
        'их количественных и стоимостных показателях, а также выявляет позиции, '
        'требующие пополнения запасов.'
    )
    doc.add_paragraph(purpose_text)
    
    doc.add_paragraph()
    
    # Получаем данные
    components = Component.objects.select_related('category').all()
    
    # Применяем фильтрацию по датам
    if date_from:
        components = components.filter(created_at__gte=date_from)
    if date_to:
        components = components.filter(created_at__lte=date_to)
    
    # Статистика
    total_components = components.count()
    low_stock = components.filter(quantity__lte=F('min_quantity')).count()
    out_of_stock = components.filter(quantity=0).count()
    total_value = sum(c.price * c.quantity for c in components)
    
    # Статистика по категориям
    categories_stats = Category.objects.annotate(
        total_items=Count('components'),
        total_quantity=Sum('components__quantity')
    ).filter(total_items__gt=0)
    
    # Общая статистика
    doc.add_heading('2. ОБЩИЕ ПОКАЗАТЕЛИ', level=1)
    
    stats_text = doc.add_paragraph()
    stats_text.add_run('Настоящий раздел содержит сводную информацию о состоянии складских запасов на момент формирования отчета.')
    
    doc.add_paragraph()
    
    stats_table = doc.add_table(rows=5, cols=2)
    stats_table.style = 'Table Grid'
    
    # Заголовок таблицы статистики
    stats_table.rows[0].cells[0].text = 'Показатель'
    stats_table.rows[0].cells[1].text = 'Значение'
    for cell in stats_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)
    
    stats_data = [
        ('Всего наименований товаров:', str(total_components)),
        ('Товаров с низким остатком:', str(low_stock)),
        ('Товаров отсутствует на складе:', str(out_of_stock)),
        ('Общая стоимость товарных запасов:', f'{total_value:,.2f} ₽')
    ]
    
    for i, (label, value) in enumerate(stats_data, start=1):
        stats_table.rows[i].cells[0].text = label
        stats_table.rows[i].cells[1].text = value
        stats_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        stats_table.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        stats_table.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Анализ по категориям
    if categories_stats.exists():
        doc.add_heading('3. АНАЛИЗ ПО КАТЕГОРИЯМ ТОВАРОВ', level=1)
        
        cat_text = doc.add_paragraph()
        cat_text.add_run(
            'Распределение товаров по категориям позволяет оценить структуру '
            'товарных запасов и выявить наиболее представленные группы товаров.'
        )
        
        doc.add_paragraph()
        
        cat_table = doc.add_table(rows=1, cols=3)
        cat_table.style = 'Table Grid'
        
        # Заголовки
        cat_headers = ['Категория', 'Количество наименований', 'Общее количество']
        for i, header in enumerate(cat_headers):
            cat_table.rows[0].cells[i].text = header
            cat_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            cat_table.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(11)
        
        for cat in categories_stats:
            row = cat_table.add_row().cells
            row[0].text = cat.name
            row[1].text = str(cat.total_items)
            row[2].text = str(cat.total_quantity or 0)
            for cell in row:
                cell.paragraphs[0].runs[0].font.size = Pt(10)
        
        doc.add_paragraph()
    
    # Детальный список товаров
    doc.add_heading('4. ДЕТАЛЬНЫЙ СПИСОК ТОВАРНЫХ ПОЗИЦИЙ', level=1)
    
    detail_text = doc.add_paragraph()
    detail_text.add_run(
        'В данном разделе представлен полный перечень товаров, находящихся на складе, '
        'с указанием артикулов, наименований, количественных и стоимостных характеристик.'
    )
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    
    # Заголовки
    headers = ['Артикул', 'Наименование', 'Категория', 'Кол-во\n(шт.)', 'Мин.\nкол-во', 'Цена\n(₽)', 'Сумма\n(₽)']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Данные
    for component in components:
        row = table.add_row().cells
        row[0].text = component.article_number
        row[1].text = component.name
        row[2].text = component.category.name if component.category else '-'
        row[3].text = str(component.quantity)
        row[4].text = str(component.min_quantity)
        row[5].text = f'{component.price:,.2f}'
        row[6].text = f'{component.price * component.quantity:,.2f}'
        
        # Форматирование
        for j, cell in enumerate(row):
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if j >= 3:  # Числовые столбцы
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Итоговая строка
    total_row = table.add_row().cells
    total_row[0].text = 'ИТОГО:'
    total_row[0].merge(total_row[5])
    total_row[0].paragraphs[0].runs[0].bold = True
    total_row[0].paragraphs[0].runs[0].font.size = Pt(10)
    total_row[6].text = f'{total_value:,.2f}'
    total_row[6].paragraphs[0].runs[0].bold = True
    total_row[6].paragraphs[0].runs[0].font.size = Pt(10)
    total_row[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()
    
    # Выводы и рекомендации
    doc.add_heading('5. ВЫВОДЫ И РЕКОМЕНДАЦИИ', level=1)
    
    conclusions = []
    
    if out_of_stock > 0:
        conclusions.append(
            f'• Отсутствуют на складе {out_of_stock} наименований товаров. '
            'Рекомендуется срочное пополнение запасов для обеспечения непрерывности работы.'
        )
    
    if low_stock > 0:
        conclusions.append(
            f'• Выявлено {low_stock} наименований товаров с низким остатком (ниже минимального уровня). '
            'Необходимо планирование заказа для предотвращения дефицита.'
        )
    
    if not conclusions:
        conclusions.append(
            '• Состояние складских запасов находится в пределах нормы. '
            'Все товарные позиции обеспечены в достаточном количестве.'
        )
    
    conclusions.append(
        f'• Общая стоимость товарных запасов составляет {total_value:,.2f} ₽. '
        'Рекомендуется регулярный мониторинг оборачиваемости для оптимизации запасов.'
    )
    
    for conclusion in conclusions:
        para = doc.add_paragraph(conclusion)
        para.paragraph_format.space_after = Pt(6)
    
    doc.add_paragraph()
    doc.add_paragraph('_' * 80)
    
    # Футер с контактами
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run(
        'ООО НПФ «Роникс-Л»\n'
        'Тел.: 8 (85595) 5-09-95 | Email: info@roniks-l.ru\n'
        f'Отчет сформирован: {timezone.now().strftime("%d.%m.%Y %H:%M")}'
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(64, 64, 64)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Сохраняем в память
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    # Возвращаем файл
    response = HttpResponse(
        output.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="inventory_report_{timezone.now().strftime("%Y%m%d")}.docx"'
    return response


@login_required
def orders_report_word(request):
    """Экспорт отчета по заявкам в Word."""
    if not (request.user.is_worker or request.user.is_superuser):
        messages.error(request, 'У вас нет доступа к этой функции.')
        from django.shortcuts import redirect
        return redirect('core:dashboard')
    
    # Создаем документ
    doc = Document()
    
    # Шапка с информацией о компании
    company_heading = doc.add_paragraph()
    company_run = company_heading.add_run('ООО НПФ «Роникс-Л»')
    company_run.bold = True
    company_run.font.size = Pt(14)
    company_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    company_details = doc.add_paragraph()
    details_run = company_details.add_run(
        'ОГРН: 1021601978177 | ИНН: 1649003743 / КПП: 164901001\n'
        '423258, Республика Татарстан, Лениногорский район,\n'
        'город Лениногорск, Крупской ул, д. 4, помещ. 3'
    )
    details_run.font.size = Pt(9)
    company_details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('_' * 80)
    
    # Заголовок отчета
    title = doc.add_heading('ОТЧЕТ ПО ОБРАБОТКЕ ЗАЯВОК КЛИЕНТОВ', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Дата формирования
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f'Дата формирования: {timezone.now().strftime("%d.%m.%Y %H:%M")}')
    date_run.italic = True
    date_run.font.size = Pt(10)
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()
    
    # Назначение отчета
    doc.add_heading('1. НАЗНАЧЕНИЕ И ОБЛАСТЬ ПРИМЕНЕНИЯ', level=1)
    purpose_text = (
        'Настоящий отчет предназначен для анализа эффективности обработки заявок клиентов '
        'сервисного центра. Отчет содержит информацию о количестве принятых, обработанных '
        'и завершенных заявок, а также финансовые показатели работы предприятия за отчетный период. '
        'Данные отчета используются для оценки загруженности сотрудников, планирования ресурсов '
        'и принятия управленческих решений.'
    )
    doc.add_paragraph(purpose_text)
    
    doc.add_paragraph()
    
    # Получаем данные
    orders = Order.objects.select_related('client', 'assigned_to').all()
    
    # Статистика
    total_orders = orders.count()
    new_orders = orders.filter(status='new').count()
    in_progress = orders.filter(status='in_progress').count()
    completed = orders.filter(status='completed').count()
    cancelled = orders.filter(status='cancelled').count()
    service_orders = orders.filter(order_type='service').count()
    component_orders = orders.filter(order_type='components').count()
    total_amount = orders.filter(status='completed').aggregate(total=Sum('total_amount'))['total'] or Decimal('0.00')
    avg_amount = total_amount / completed if completed > 0 else Decimal('0.00')
    
    # Общая статистика
    doc.add_heading('2. СВОДНЫЕ ПОКАЗАТЕЛИ', level=1)
    
    stats_intro = doc.add_paragraph()
    stats_intro.add_run(
        'В данном разделе представлена сводная информация о заявках, '
        'сгруппированная по статусам обработки и типам оказываемых услуг.'
    )
    
    doc.add_paragraph()
    
    stats_table = doc.add_table(rows=8, cols=2)
    stats_table.style = 'Table Grid'
    
    # Заголовок таблицы
    stats_table.rows[0].cells[0].text = 'Показатель'
    stats_table.rows[0].cells[1].text = 'Значение'
    for cell in stats_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)
    
    stats_data = [
        ('Всего заявок за период:', str(total_orders)),
        ('Новые заявки (ожидают обработки):', str(new_orders)),
        ('Заявки в работе:', str(in_progress)),
        ('Завершенные заявки:', str(completed)),
        ('Отмененные заявки:', str(cancelled)),
        ('Заявки на сервисное обслуживание:', str(service_orders)),
        ('Заявки на продажу комплектующих:', str(component_orders))
    ]
    
    for i, (label, value) in enumerate(stats_data, start=1):
        stats_table.rows[i].cells[0].text = label
        stats_table.rows[i].cells[1].text = value
        stats_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        stats_table.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        stats_table.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Финансовые показатели
    doc.add_heading('3. ФИНАНСОВЫЕ РЕЗУЛЬТАТЫ', level=1)
    
    fin_text = doc.add_paragraph()
    fin_text.add_run(
        'Раздел содержит финансовые показатели по завершенным заявкам, '
        'на основании которых может быть проведен анализ рентабельности работы сервисного центра.'
    )
    
    doc.add_paragraph()
    
    fin_table = doc.add_table(rows=3, cols=2)
    fin_table.style = 'Table Grid'
    
    fin_table.rows[0].cells[0].text = 'Финансовый показатель'
    fin_table.rows[0].cells[1].text = 'Сумма (₽)'
    for cell in fin_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)
    
    fin_data = [
        ('Общая выручка по завершенным заявкам:', f'{total_amount:,.2f}'),
        ('Средний чек (на одну завершенную заявку):', f'{avg_amount:,.2f}')
    ]
    
    for i, (label, value) in enumerate(fin_data, start=1):
        fin_table.rows[i].cells[0].text = label
        fin_table.rows[i].cells[1].text = value
        fin_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        fin_table.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        fin_table.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        fin_table.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()
    
    # Детальный список заявок
    doc.add_heading('4. ДЕТАЛЬНЫЙ ПЕРЕЧЕНЬ ЗАЯВОК', level=1)
    
    detail_intro = doc.add_paragraph()
    detail_intro.add_run(
        'Ниже представлен полный список всех заявок за отчетный период с указанием '
        'клиентов, типов услуг, статусов обработки, назначенных исполнителей и финансовых показателей.'
    )
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    
    # Заголовки
    headers = ['№\nзаявки', 'Дата\nсоздания', 'Клиент', 'Тип\nуслуги', 'Устройство', 'Статус', 'Исполнитель', 'Сумма\n(₽)']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    status_dict = dict(Order.STATUS_CHOICES)
    type_dict = dict(Order.ORDER_TYPE_CHOICES)
    
    # Данные
    for order in orders:
        row = table.add_row().cells
        row[0].text = str(order.id)
        row[1].text = order.created_at.strftime('%d.%m.%Y\n%H:%M')
        row[2].text = order.client.get_full_name() or order.client.username
        row[3].text = type_dict.get(order.order_type, order.order_type)
        # Показываем первые 40 символов описания или прочерк
        desc_text = order.description[:40] + '...' if len(order.description) > 40 else order.description
        row[4].text = desc_text if desc_text else '-'
        row[5].text = status_dict.get(order.status, order.status)
        row[6].text = order.assigned_to.get_full_name() if order.assigned_to else 'Не назначен'
        row[7].text = f'{order.total_amount:,.2f}' if order.total_amount else '0.00'
        
        # Форматирование
        for j, cell in enumerate(row):
            cell.paragraphs[0].runs[0].font.size = Pt(8)
            if j in [0, 7]:  # Числовые столбцы
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Итоговая строка
    total_row = table.add_row().cells
    total_row[0].text = 'ИТОГО по завершенным:'
    total_row[0].merge(total_row[6])
    total_row[0].paragraphs[0].runs[0].bold = True
    total_row[0].paragraphs[0].runs[0].font.size = Pt(9)
    completed_total = orders.filter(status='completed').aggregate(total=Sum('total_amount'))['total'] or Decimal('0.00')
    total_row[7].text = f'{completed_total:,.2f}'
    total_row[7].paragraphs[0].runs[0].bold = True
    total_row[7].paragraphs[0].runs[0].font.size = Pt(9)
    total_row[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()
    
    # Анализ и выводы
    doc.add_heading('5. АНАЛИЗ И РЕКОМЕНДАЦИИ', level=1)
    
    conclusions = []
    
    completion_rate = (completed / total_orders * 100) if total_orders > 0 else 0
    conclusions.append(
        f'• Процент завершенных заявок составляет {completion_rate:.1f}% от общего числа. '
        f'Завершено {completed} из {total_orders} заявок.'
    )
    
    if new_orders > 0:
        conclusions.append(
            f'• В настоящий момент {new_orders} заявок ожидают назначения исполнителя. '
            'Рекомендуется распределить нагрузку между доступными сотрудниками.'
        )
    
    if cancelled > 0:
        cancellation_rate = (cancelled / total_orders * 100) if total_orders > 0 else 0
        conclusions.append(
            f'• Отменено {cancelled} заявок ({cancellation_rate:.1f}% от общего числа). '
            'Рекомендуется провести анализ причин отмены для улучшения качества обслуживания.'
        )
    
    if completed > 0:
        conclusions.append(
            f'• Средний чек по завершенным заявкам составляет {avg_amount:,.2f} ₽. '
            'Данный показатель может использоваться для планирования выручки.'
        )
    
    conclusions.append(
        f'• Общая выручка по завершенным заявкам за период составила {total_amount:,.2f} ₽. '
        'Рекомендуется регулярный мониторинг динамики для своевременного выявления трендов.'
    )
    
    for conclusion in conclusions:
        para = doc.add_paragraph(conclusion)
        para.paragraph_format.space_after = Pt(6)
    
    doc.add_paragraph()
    doc.add_paragraph('_' * 80)
    
    # Футер с контактами
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run(
        'ООО НПФ «Роникс-Л»\n'
        'Тел.: 8 (85595) 5-09-95 | Email: info@roniks-l.ru\n'
        f'Отчет сформирован: {timezone.now().strftime("%d.%m.%Y %H:%M")}'
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(64, 64, 64)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Сохраняем в память
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    # Возвращаем файл
    response = HttpResponse(
        output.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="orders_report_{timezone.now().strftime("%Y%m%d")}.docx"'
    return response


# ==================== IMPROVED PDF EXPORTS ====================

@login_required
def inventory_report_pdf_improved(request):
    """Улучшенный PDF отчет по складу."""
    if not (request.user.is_worker or request.user.is_superuser):
        messages.error(request, 'У вас нет доступа к этой функции.')
        from django.shortcuts import redirect
        return redirect('core:dashboard')
    
    # Получаем параметры фильтрации
    date_from = request.GET.get('date_from')
    date_to = request.GET.get('date_to')
    
    components = Component.objects.select_related('category').all()
    
    # Применяем фильтрацию по датам
    if date_from:
        components = components.filter(created_at__gte=date_from)
    if date_to:
        components = components.filter(created_at__lte=date_to)
    
    # Добавляем вычисленную сумму каждому компоненту
    components_with_total = []
    for c in components:
        component_data = {
            'obj': c,
            'total': c.price * c.quantity
        }
        components_with_total.append(component_data)
    
    # Статистика
    total_components = len(components_with_total)
    low_stock = components.filter(quantity__lte=F('min_quantity')).count()
    out_of_stock = components.filter(quantity=0).count()
    total_value = sum(item['total'] for item in components_with_total)
    
    html = render_to_string('reports/inventory_report_pdf_improved.html', {
        'components': components_with_total,
        'total_components': total_components,
        'low_stock': low_stock,
        'out_of_stock': out_of_stock,
        'total_value': total_value,
        'date': timezone.now(),
    })
    
    # Генерируем PDF (если WeasyPrint доступен)
    if WEASYPRINT_AVAILABLE:
        try:
            pdf_file = HTML(string=html).write_pdf()
            response = HttpResponse(pdf_file, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="inventory_report_{timezone.now().strftime("%Y%m%d")}.pdf"'
            return response
        except Exception as e:
            messages.warning(request, f'Ошибка генерации PDF: {e}. Показываем HTML версию.')
    
    # Fallback - возвращаем HTML версию для печати
    response = HttpResponse(html, content_type='text/html')
    return response


@login_required
def orders_report_pdf_improved(request):
    """Улучшенный PDF отчет по заявкам."""
    if not (request.user.is_worker or request.user.is_superuser):
        messages.error(request, 'У вас нет доступа к этой функции.')
        from django.shortcuts import redirect
        return redirect('core:dashboard')
    
    # Получаем параметры фильтрации
    date_from = request.GET.get('date_from')
    date_to = request.GET.get('date_to')
    
    orders = Order.objects.select_related('client', 'assigned_to').all()
    
    # Применяем фильтрацию по датам
    if date_from:
        orders = orders.filter(created_at__gte=date_from)
    if date_to:
        orders = orders.filter(created_at__lte=date_to)
    
    # Статистика
    total_orders = orders.count()
    new_orders = orders.filter(status='new').count()
    in_progress = orders.filter(status='in_progress').count()
    completed = orders.filter(status='completed').count()
    cancelled = orders.filter(status='cancelled').count()
    total_amount = orders.filter(status='completed').aggregate(total=Sum('total_amount'))['total'] or Decimal('0.00')
    
    html = render_to_string('reports/orders_report_pdf_improved.html', {
        'orders': orders,
        'total_orders': total_orders,
        'new_orders': new_orders,
        'in_progress': in_progress,
        'completed': completed,
        'cancelled': cancelled,
        'total_amount': total_amount,
        'date': timezone.now(),
    })
    
    # Генерируем PDF (если WeasyPrint доступен)
    if WEASYPRINT_AVAILABLE:
        try:
            pdf_file = HTML(string=html).write_pdf()
            response = HttpResponse(pdf_file, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="orders_report_{timezone.now().strftime("%Y%m%d")}.pdf"'
            return response
        except Exception as e:
            messages.warning(request, f'Ошибка генерации PDF: {e}. Показываем HTML версию.')
    
    # Fallback - возвращаем HTML версию для печати
    response = HttpResponse(html, content_type='text/html')
    return response
# Временный файл с функцией orders_report_word для добавления в views.py

@login_required
def orders_report_word(request):
    """Экспорт отчета по заявкам в Word."""
    if not (request.user.is_worker or request.user.is_superuser):
        messages.error(request, 'У вас нет доступа к этой функции.')
        from django.shortcuts import redirect
        return redirect('core:dashboard')
    
    # Получаем параметры фильтрации
    date_from = request.GET.get('date_from')
    date_to = request.GET.get('date_to')
    
    # Создаем документ
    doc = Document()
    
    # Шапка с информацией о компании
    company_heading = doc.add_paragraph()
    company_run = company_heading.add_run('ООО НПФ «Роникс-Л»')
    company_run.bold = True
    company_run.font.size = Pt(14)
    company_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    company_details = doc.add_paragraph()
    details_run = company_details.add_run(
        'ОГРН: 1021601978177 | ИНН: 1649003743 / КПП: 164901001\n'
        '423258, Республика Татарстан, Лениногорский район,\n'
        'город Лениногорск, Крупской ул, д. 4, помещ. 3'
    )
    details_run.font.size = Pt(9)
    company_details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('_' * 80)
    
    # Заголовок отчета
    title = doc.add_heading('ОТЧЕТ ПО ЗАЯВКАМ', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Период отчета
    if date_from or date_to:
        period_para = doc.add_paragraph()
        period_text = 'Период: '
        if date_from:
            period_text += f'с {date_from} '
        if date_to:
            period_text += f'по {date_to}'
        period_run = period_para.add_run(period_text)
        period_run.bold = True
        period_run.font.size = Pt(11)
        period_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Дата формирования
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f'Дата формирования: {timezone.now().strftime("%d.%m.%Y %H:%M")}')
    date_run.italic = True
    date_run.font.size = Pt(10)
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()
    
    # Получаем данные
    orders = Order.objects.select_related('client', 'assigned_to').all()
    
    # Применяем фильтрацию по датам
    if date_from:
        orders = orders.filter(created_at__gte=date_from)
    if date_to:
        orders = orders.filter(created_at__lte=date_to)
    
    # Статистика
    total_orders = orders.count()
    new_orders = orders.filter(status='new').count()
    in_progress_orders = orders.filter(status='in_progress').count()
    completed_orders = orders.filter(status='completed').count()
    cancelled_orders = orders.filter(status='cancelled').count()
    
    # Общая статистика
    doc.add_heading('1. ОБЩИЕ ПОКАЗАТЕЛИ', level=1)
    
    stats_table = doc.add_table(rows=6, cols=2)
    stats_table.style = 'Table Grid'
    
    # Заголовок таблицы
    stats_table.rows[0].cells[0].text = 'Показатель'
    stats_table.rows[0].cells[1].text = 'Значение'
    for cell in stats_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)
    
    stats_data = [
        ('Всего заявок:', str(total_orders)),
        ('Новые заявки:', str(new_orders)),
        ('В работе:', str(in_progress_orders)),
        ('Завершенные:', str(completed_orders)),
        ('Отмененные:', str(cancelled_orders))
    ]
    
    for i, (label, value) in enumerate(stats_data, start=1):
        stats_table.rows[i].cells[0].text = label
        stats_table.rows[i].cells[1].text = value
        stats_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        stats_table.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        stats_table.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Детальный список заявок
    doc.add_heading('2. ДЕТАЛЬНЫЙ СПИСОК ЗАЯВОК', level=1)
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    
    # Заголовки
    headers = ['№', 'Дата', 'Клиент', 'Тип', 'Статус', 'Исполнитель', 'Сумма (₽)']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Словари для отображения
    status_dict = dict(Order.STATUS_CHOICES)
    type_dict = dict(Order.ORDER_TYPE_CHOICES)
    
    # Данные
    total_sum = Decimal('0.00')
    for order in orders:
        row = table.add_row().cells
        row[0].text = str(order.id)
        row[1].text = order.created_at.strftime('%d.%m.%Y %H:%M')
        row[2].text = order.client.get_full_name() or order.client.username
        row[3].text = type_dict.get(order.order_type, order.order_type)
        row[4].text = status_dict.get(order.status, order.status)
        row[5].text = order.assigned_to.get_full_name() if order.assigned_to else 'Не назначен'
        row[6].text = f'{order.total_amount:,.2f}'
        
        if order.status == 'completed':
            total_sum += order.total_amount
        
        # Форматирование
        for j, cell in enumerate(row):
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if j == 6:  # Сумма
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Итоговая строка
    total_row = table.add_row().cells
    total_row[0].text = 'ИТОГО (завершенные):'
    total_row[0].merge(total_row[5])
    total_row[0].paragraphs[0].runs[0].bold = True
    total_row[0].paragraphs[0].runs[0].font.size = Pt(10)
    total_row[6].text = f'{total_sum:,.2f}'
    total_row[6].paragraphs[0].runs[0].bold = True
    total_row[6].paragraphs[0].runs[0].font.size = Pt(10)
    total_row[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()
    doc.add_paragraph('_' * 80)
    
    # Футер
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run(
        'ООО НПФ «Роникс-Л»\n'
        'Тел.: 8 (85595) 5-09-95 | Email: info@roniks-l.ru\n'
        f'Отчет сформирован: {timezone.now().strftime("%d.%m.%Y %H:%M")}'
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(64, 64, 64)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Сохраняем в память
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    # Возвращаем файл
    response = HttpResponse(
        output.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="orders_report_{timezone.now().strftime("%Y%m%d")}.docx"'
    return response
